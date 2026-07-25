"""DOCX rendering — structure, tables, watermark. Pure function, no live anything."""

from __future__ import annotations

import io

from docx import Document

from app.docx_export import render

CTX = {
    "tender": {"title": "e-Office Software Implementation",
               "tender_number": "MAHA/IT/2026/4415", "authority": "MahaIT"},
    "bidder": {"name": "Meridian Infotech", "cin": "U72900MH2015PTC1", "gst": "27AAA1Z"},
    "generated_on": "2026-07-25",
    "approved": False,
    "sections": [
        {"key": "understanding", "heading": "Form 7(b): Understanding", "kind": "narrative",
         "status": "drafted", "approved": False,
         "body_md": "### Landscape\n\nThe department requires workflow automation.\n\n"
                    "### Scope\n\n- Digitisation\n- Workflow\n"},
        {"key": "project_citations", "heading": "Form 6: Project Citations", "kind": "assembled",
         "status": "drafted", "approved": True,
         "body_md": "| Project | Value |\n| --- | --- |\n| HMIS | ₹2.40 Cr |\n| e-Office | ₹3.80 Cr |"},
    ],
}


def _doc(ctx=None):
    return Document(io.BytesIO(render(ctx or CTX)))


def _text(doc):
    return "\n".join(p.text for p in doc.paragraphs)


def test_render_returns_a_real_docx():
    blob = render(CTX)
    assert blob[:2] == b"PK"  # OOXML is a zip
    assert len(blob) > 5_000


def test_cover_carries_tender_identity():
    t = _text(_doc())
    assert "TECHNICAL PROPOSAL" in t
    assert "MAHA/IT/2026/4415" in t
    assert "MahaIT" in t
    assert "U72900MH2015PTC1" in t


def test_every_section_heading_is_rendered():
    headings = [p.text for p in _doc().paragraphs if p.style.name.startswith("Heading")]
    assert any("Form 7(b): Understanding" in h for h in headings)
    assert any("Form 6: Project Citations" in h for h in headings)


def test_subsection_headings_become_level_two():
    doc = _doc()
    h2 = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
    assert "Landscape" in h2 and "Scope" in h2


def test_markdown_table_becomes_a_real_table_with_a_repeating_header():
    doc = _doc()
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert table.rows[0].cells[0].text == "Project"
    assert table.rows[1].cells[1].text == "₹2.40 Cr"
    # tblHeader makes Word repeat the header across page breaks — evaluators read these
    # matrices across many pages.
    assert "tblHeader" in table.rows[0]._tr.xml


def test_bullets_render_as_list_items():
    doc = _doc()
    bullets = [p.text for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert "Digitisation" in bullets


def test_unapproved_narrative_carries_the_watermark():
    t = _text(_doc())
    assert "AI DRAFT — NOT APPROVED" in t
    assert "must not be submitted in this state" in t


def test_notice_page_names_only_unapproved_narrative_sections():
    t = _text(_doc())
    assert "Form 7(b): Understanding" in t
    # the assembled section is approved and must not appear in the awaiting-approval list
    assert t.count("Form 6: Project Citations") == 1


def test_page_header_marks_every_page_when_unapproved():
    doc = _doc()
    header = "\n".join(p.text for p in doc.sections[0].header.paragraphs)
    assert "NOT FOR SUBMISSION" in header


def test_approved_document_has_no_watermark():
    ctx = {**CTX, "approved": True,
           "sections": [{**s, "approved": True} for s in CTX["sections"]]}
    doc = _doc(ctx)
    t = _text(doc)
    header = "\n".join(p.text for p in doc.sections[0].header.paragraphs)
    assert "AI DRAFT" not in t
    assert "NOT FOR SUBMISSION" not in header


def test_table_of_contents_field_is_present():
    doc = _doc()
    assert "TOC" in doc.element.xml


def test_empty_sections_do_not_crash():
    blob = render({"tender": {}, "bidder": {}, "sections": []})
    assert blob[:2] == b"PK"


def test_section_with_empty_body_is_still_headed():
    doc = _doc({**CTX, "sections": [
        {"key": "risk", "heading": "Risk", "kind": "narrative", "status": "placeholder",
         "approved": False, "body_md": ""}
    ]})
    assert any("Risk" in p.text for p in doc.paragraphs if p.style.name.startswith("Heading"))
