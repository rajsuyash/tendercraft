"""DOCX rendering of the long-form proposal — the actual deliverable.

Pure function over a context dict: no DB, no HTTP, so it unit-tests without a live
anything. The export GATE runs before this is ever called (app/proposal_routes.py), so no
bytes exist for a proposal that may not be exported.

Markdown subset understood: "### " subheadings, "| " tables, "- " bullets, blank-line
paragraphs. That is exactly what the assemblers and the section drafter emit — a full
markdown parser would be a dependency bought for nothing.
"""

from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# design/tokens.json colour.warning — the AI-draft mark reads as a caution, not decoration.
_WARNING = RGBColor(0xB4, 0x53, 0x09)
_MUTED = RGBColor(0x64, 0x74, 0x8B)

_AI_TAG = "  [AI DRAFT — NOT APPROVED]"
_PAGE_MARK = "AI-GENERATED DRAFT — NOT FOR SUBMISSION"


def _style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    # ponytail: Calibri/Cambria, the Word-native safe pair. The brand fonts (Lexend/Inter)
    # are not installed on an evaluator's machine and would silently substitute; embedding
    # them needs a font pack (PH2).
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    for name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)):
        st = doc.styles[name]
        st.font.name = "Cambria"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)  # colour.primary


def _repeat_header_row(row) -> None:
    """Mark a table row as a header so Word repeats it across page breaks.

    Evaluators read these matrices across many pages; a table whose header appears once is
    genuinely harder to score.
    """
    tr_pr = row._tr.get_or_add_trPr()
    el = tr_pr.makeelement(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblHeader", {}
    )
    tr_pr.append(el)


def _add_table(doc: Document, lines: list[str]) -> None:
    rows = [
        [c.strip() for c in ln.strip().strip("|").split("|")]
        for ln in lines
        if ln.strip() and not re.match(r"^\|[\s\-|]+\|$", ln.strip())
    ]
    if not rows:
        return
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=width)
    table.style = "Table Grid"
    for i, cells in enumerate(rows):
        cells = cells + [""] * (width - len(cells))
        wr = table.add_row()
        for j, text in enumerate(cells):
            cell = wr.cells[j]
            cell.text = text
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
        if i == 0:
            _repeat_header_row(wr)
    doc.add_paragraph()


def _add_markdown(doc: Document, body: str) -> None:
    """Render the emitted markdown subset. Unknown syntax falls through as plain text."""
    buf: list[str] = []

    def flush_table():
        if buf:
            _add_table(doc, buf)
            buf.clear()

    for raw in (body or "").splitlines():
        line = raw.rstrip()
        if line.lstrip().startswith("|"):
            buf.append(line)
            continue
        flush_table()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            p = doc.add_paragraph(_strip_emphasis(line))
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    flush_table()


def _strip_emphasis(text: str) -> str:
    """Flatten **bold** / _italic_ markers — we render structure, not inline styling."""
    return re.sub(r"\*\*(.+?)\*\*", r"\1", re.sub(r"(?<!\w)_(.+?)_(?!\w)", r"\1", text))


def _cover(doc: Document, ctx: dict[str, Any]) -> None:
    tender = ctx.get("tender") or {}
    bidder = ctx.get("bidder") or {}

    for _ in range(4):
        doc.add_paragraph()
    h = doc.add_heading("TECHNICAL PROPOSAL", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for text, size, bold in (
        (tender.get("title") or "Untitled tender", 15, True),
        (f"Tender No. {tender.get('tender_number') or '—'}", 12, False),
        (f"Issued by: {tender.get('authority') or '—'}", 12, False),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold

    doc.add_paragraph()
    for text in (
        f"Submitted by: {bidder.get('name') or 'Bidder'}",
        f"CIN: {bidder.get('cin') or '—'}    GST: {bidder.get('gst') or '—'}",
        f"Date: {ctx.get('generated_on') or '—'}",
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(text).font.size = Pt(11)


def _draft_notice(doc: Document, unapproved: list[str]) -> None:
    doc.add_page_break()
    doc.add_heading("Notice — AI-assisted draft", level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        "This document contains AI-drafted narrative that has NOT been approved by an "
        "authorised reviewer. It must not be submitted in this state."
    )
    run.bold = True
    run.font.color.rgb = _WARNING
    doc.add_paragraph("Sections awaiting human approval:")
    for key in unapproved:
        doc.add_paragraph(key, style="List Bullet")


def _toc(doc: Document) -> None:
    """A TOC field. Word populates it on open — we do not hand-build page numbers.

    ponytail: field code, not a rendered TOC. A pre-paginated one needs a headless Word or
    LibreOffice pass, which is a deployment problem, not a rendering one.
    """
    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    fld = p.add_run()._r.makeelement(
        f"{ns}fldSimple", {f"{ns}instr": 'TOC \\o "1-2" \\h'}
    )
    p._p.append(fld)
    doc.add_paragraph(
        "(Right-click and select 'Update Field' in Word to populate.)"
    ).runs[0].font.color.rgb = _MUTED


def _page_watermark(doc: Document) -> None:
    for section in doc.sections:
        para = section.header.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(_PAGE_MARK)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = _WARNING


def render(ctx: dict[str, Any]) -> bytes:
    """Render the proposal to .docx bytes.

    ctx: {tender, bidder, generated_on, approved: bool, sections: [{key, heading, body_md,
    kind, status, approved}]}
    """
    sections = ctx.get("sections") or []
    fully_approved = bool(ctx.get("approved"))
    unapproved = [
        s.get("heading") or s.get("key")
        for s in sections
        if s.get("kind") == "narrative" and not s.get("approved")
    ]

    doc = Document()
    _style_document(doc)
    _cover(doc, ctx)

    if not fully_approved:
        _draft_notice(doc, unapproved)
        _page_watermark(doc)

    _toc(doc)

    for s in sections:
        doc.add_page_break()
        heading = doc.add_heading(s.get("heading") or s.get("key") or "Section", level=1)
        # Per-section mark: an evaluator flipping to one section must see its status there,
        # not only on the notice page.
        if s.get("kind") == "narrative" and not s.get("approved"):
            run = heading.add_run(_AI_TAG)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = _WARNING
        _add_markdown(doc, s.get("body_md") or "")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


__all__ = ["render", "WD_SECTION"]
